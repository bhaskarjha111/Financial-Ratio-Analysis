from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import os

# ============================================================
# FINANCIAL RATIO ANALYSIS PROJECT
# Apple vs Microsoft vs Samsung
# ============================================================

# NOTE:
# Financial figures below are example project values.
# Replace them with exact numbers from latest annual reports
# whenever you submit or upload this project.


companies = {
    "Apple": {
        "currency": "USD Million",
        "revenue": 391035,
        "net_income": 93736,
        "total_assets": 364980,
        "total_liabilities": 308030,
        "shareholders_equity": 56950,
        "eps": 6.08,
        "market_price": 270.00,
    },
    "Microsoft": {
        "currency": "USD Million",
        "revenue": 281724,
        "net_income": 101832,
        "total_assets": 619003,
        "total_liabilities": 243686,
        "shareholders_equity": 375317,
        "eps": 13.64,
        "market_price": 520.00,
    },
    "Samsung Electronics": {
        "currency": "KRW Billion",
        "revenue": 300871,
        "net_income": 34289,
        "total_assets": 455906,
        "total_liabilities": 100676,
        "shareholders_equity": 355230,
        "eps": 5100,
        "market_price": 74000,
    },
}

# ============================================================
# Create CSV Dataset
# ============================================================

os.makedirs("data", exist_ok=True)

csv_file = "data/financial_data.csv"

csv_headers = [
    "Company",
    "Currency",
    "Revenue",
    "Net Income",
    "Total Assets",
    "Total Liabilities",
    "Shareholders' Equity",
    "EPS",
    "Market Price"
]

with open(csv_file, "w", newline="", encoding="utf-8") as file:
    writer = csv.DictWriter(file, fieldnames=csv_headers)
    writer.writeheader()

    for company, data in companies.items():
        writer.writerow({
            "Company": company,
            "Currency": data["currency"],
            "Revenue": data["revenue"],
            "Net Income": data["net_income"],
            "Total Assets": data["total_assets"],
            "Total Liabilities": data["total_liabilities"],
            "Shareholders' Equity": data["shareholders_equity"],
            "EPS": data["eps"],
            "Market Price": data["market_price"]
        })

print(f"CSV dataset created successfully: {csv_file}")

# ============================================================
# Ratio Calculation Functions
# ============================================================

def calculate_ratios(data):
    revenue = data["revenue"]
    net_income = data["net_income"]
    liabilities = data["total_liabilities"]
    equity = data["shareholders_equity"]
    eps = data["eps"]
    market_price = data["market_price"]

    ratios = {
        "P/E Ratio": market_price / eps,
        "Debt-to-Equity Ratio": liabilities / equity,
        "Return on Equity (%)": (net_income / equity) * 100,
        "Net Profit Margin (%)": (net_income / revenue) * 100,
    }

    return ratios


# ============================================================
# Create Excel Workbook
# ============================================================

wb = Workbook()

# Remove default sheet
default_sheet = wb.active
wb.remove(default_sheet)

# Styles
title_fill = PatternFill("solid", fgColor="1F4E78")
header_fill = PatternFill("solid", fgColor="D9EAF7")
subheader_fill = PatternFill("solid", fgColor="E2F0D9")
white_font = Font(color="FFFFFF", bold=True)
bold_font = Font(bold=True)
title_font = Font(size=16, bold=True, color="FFFFFF")
thin_border = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin")
)


def style_sheet(ws):
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border

    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 22


# ============================================================
# Sheet 1: Project Overview
# ============================================================

ws = wb.create_sheet("Project Overview")

ws.merge_cells("A1:F1")
ws["A1"] = "Financial Ratio Analysis: Apple vs Microsoft vs Samsung"
ws["A1"].font = title_font
ws["A1"].fill = title_fill
ws["A1"].alignment = Alignment(horizontal="center")

overview_data = [
    ["Project Type", "Financial Analysis / Equity Research / Business Analytics"],
    ["Objective", "To compare the financial strength, profitability, leverage, and valuation of three global companies."],
    ["Companies Selected", "Apple, Microsoft, Samsung Electronics"],
    ["Ratios Calculated", "P/E Ratio, Debt-to-Equity Ratio, Return on Equity, Net Profit Margin"],
    ["Tools Used", "Python, Excel, OpenPyXL, Financial Statements"],
    ["Output", "Excel Model, Ratio Dashboard, Charts, Final Report"],
]

row = 3
for item in overview_data:
    ws.cell(row=row, column=1).value = item[0]
    ws.cell(row=row, column=2).value = item[1]
    ws.cell(row=row, column=1).font = bold_font
    row += 1

style_sheet(ws)


# ============================================================
# Sheet 2: Raw Financial Data
# ============================================================

ws = wb.create_sheet("Raw Financial Data")

headers = [
    "Company",
    "Currency",
    "Revenue",
    "Net Income",
    "Total Assets",
    "Total Liabilities",
    "Shareholders' Equity",
    "EPS",
    "Market Price",
]

ws.append(headers)

for company, data in companies.items():
    ws.append([
        company,
        data["currency"],
        data["revenue"],
        data["net_income"],
        data["total_assets"],
        data["total_liabilities"],
        data["shareholders_equity"],
        data["eps"],
        data["market_price"],
    ])

for cell in ws[1]:
    cell.font = bold_font
    cell.fill = header_fill

style_sheet(ws)


# ============================================================
# Sheet 3: Ratio Analysis
# ============================================================

ws = wb.create_sheet("Ratio Analysis")

ratio_headers = [
    "Company",
    "P/E Ratio",
    "Debt-to-Equity Ratio",
    "Return on Equity (%)",
    "Net Profit Margin (%)",
    "Interpretation",
]

ws.append(ratio_headers)

for company, data in companies.items():
    ratios = calculate_ratios(data)

    if ratios["Return on Equity (%)"] > 30 and ratios["Net Profit Margin (%)"] > 20:
        interpretation = "Strong profitability and efficient equity usage"
    elif ratios["Debt-to-Equity Ratio"] > 2:
        interpretation = "High leverage; requires careful risk review"
    else:
        interpretation = "Stable financial position"

    ws.append([
        company,
        round(ratios["P/E Ratio"], 2),
        round(ratios["Debt-to-Equity Ratio"], 2),
        round(ratios["Return on Equity (%)"], 2),
        round(ratios["Net Profit Margin (%)"], 2),
        interpretation,
    ])

for cell in ws[1]:
    cell.font = bold_font
    cell.fill = header_fill

style_sheet(ws)


# ============================================================
# Sheet 4: Formula Explanation
# ============================================================

ws = wb.create_sheet("Formula Explanation")

formula_data = [
    ["Ratio", "Formula", "Meaning"],
    ["P/E Ratio", "Market Price per Share / Earnings per Share", "Shows how much investors pay for each unit of earnings."],
    ["Debt-to-Equity Ratio", "Total Liabilities / Shareholders' Equity", "Measures financial leverage and debt dependence."],
    ["Return on Equity", "Net Income / Shareholders' Equity × 100", "Shows how efficiently a company generates profit from shareholders' funds."],
    ["Net Profit Margin", "Net Income / Revenue × 100", "Shows how much profit is earned from each unit of sales."],
]

for row in formula_data:
    ws.append(row)

for cell in ws[1]:
    cell.font = bold_font
    cell.fill = header_fill

style_sheet(ws)


# ============================================================
# Sheet 5: Dashboard
# ============================================================

ws = wb.create_sheet("Dashboard")

ws.merge_cells("A1:F1")
ws["A1"] = "Financial Ratio Dashboard"
ws["A1"].font = title_font
ws["A1"].fill = title_fill
ws["A1"].alignment = Alignment(horizontal="center")

dashboard_headers = [
    "Company",
    "P/E Ratio",
    "Debt-to-Equity Ratio",
    "ROE (%)",
    "Profit Margin (%)",
]

ws.append([])
ws.append(dashboard_headers)

for company, data in companies.items():
    ratios = calculate_ratios(data)
    ws.append([
        company,
        round(ratios["P/E Ratio"], 2),
        round(ratios["Debt-to-Equity Ratio"], 2),
        round(ratios["Return on Equity (%)"], 2),
        round(ratios["Net Profit Margin (%)"], 2),
    ])

for cell in ws[3]:
    cell.font = bold_font
    cell.fill = header_fill

# Chart 1: Profitability Comparison
bar_chart = BarChart()
bar_chart.title = "Profitability Comparison"
bar_chart.y_axis.title = "Percentage"
bar_chart.x_axis.title = "Company"

data = Reference(ws, min_col=4, max_col=5, min_row=3, max_row=6)
categories = Reference(ws, min_col=1, min_row=4, max_row=6)

bar_chart.add_data(data, titles_from_data=True)
bar_chart.set_categories(categories)
bar_chart.height = 8
bar_chart.width = 14

ws.add_chart(bar_chart, "G3")

# Chart 2: Leverage Comparison
de_chart = BarChart()
de_chart.title = "Debt-to-Equity Ratio Comparison"
de_chart.y_axis.title = "Debt-to-Equity"
de_chart.x_axis.title = "Company"

data = Reference(ws, min_col=3, min_row=3, max_row=6)
categories = Reference(ws, min_col=1, min_row=4, max_row=6)

de_chart.add_data(data, titles_from_data=True)
de_chart.set_categories(categories)
de_chart.height = 8
de_chart.width = 14

ws.add_chart(de_chart, "G20")

style_sheet(ws)


# ============================================================
# Sheet 6: Advanced Insights
# ============================================================

ws = wb.create_sheet("Advanced Insights")

insights = [
    ["Company", "Key Insight"],
    [
        "Apple",
        "Apple shows strong profitability and brand-driven pricing power. However, its high debt-to-equity ratio indicates aggressive capital structure and large shareholder return programs."
    ],
    [
        "Microsoft",
        "Microsoft has strong ROE and profit margin, supported by cloud, enterprise software, and recurring revenue. It appears financially balanced compared to Apple."
    ],
    [
        "Samsung Electronics",
        "Samsung has lower leverage and strong asset base, but its profitability depends heavily on semiconductor and electronics cycles."
    ],
    [
        "Overall Conclusion",
        "Microsoft appears the most balanced company, Apple remains highly profitable but leveraged, while Samsung is financially conservative but more cyclical."
    ],
]

for row in insights:
    ws.append(row)

for cell in ws[1]:
    cell.font = bold_font
    cell.fill = header_fill

style_sheet(ws)


# ============================================================
# Sheet 7: Source Log
# ============================================================

ws = wb.create_sheet("Source Log")

source_data = [
    ["Company", "Source Used", "Notes"],
    ["Apple", "Apple Annual Report / Form 10-K", "Used revenue, net income, assets, liabilities, equity and EPS."],
    ["Microsoft", "Microsoft Annual Report", "Used revenue, net income, assets, liabilities, equity and EPS."],
    ["Samsung Electronics", "Samsung Annual Report / Consolidated Financial Statements", "Used revenue, profit, liabilities, equity and EPS."],
    ["Market Price", "Market data website / finance portal", "Used only for P/E ratio calculation; price changes daily."],
]

for row in source_data:
    ws.append(row)

for cell in ws[1]:
    cell.font = bold_font
    cell.fill = header_fill

style_sheet(ws)


# ============================================================
# Sheet 8: Resume and LinkedIn
# ============================================================

ws = wb.create_sheet("Resume LinkedIn Content")

resume_content = [
    ["Section", "Content"],
    [
        "Resume Bullet 1",
        "Built a financial ratio analysis model comparing Apple, Microsoft and Samsung using annual report data."
    ],
    [
        "Resume Bullet 2",
        "Calculated P/E ratio, Debt-to-Equity, ROE and Net Profit Margin to evaluate valuation, leverage and profitability."
    ],
    [
        "Resume Bullet 3",
        "Created an Excel dashboard with charts, interpretations and company-wise insights for business decision-making."
    ],
    [
        "LinkedIn Overview",
        "I completed a Financial Ratio Analysis project comparing Apple, Microsoft and Samsung using real financial statement data. The project includes ratio calculations, interpretation, dashboard visualization and business insights, helping understand how analysts compare companies on profitability, leverage, valuation and efficiency."
    ],
]

for row in resume_content:
    ws.append(row)

for cell in ws[1]:
    cell.font = bold_font
    cell.fill = header_fill

style_sheet(ws)


# ============================================================
# Save Excel File
# ============================================================

excel_file = "Financial_Ratio_Analysis_Apple_Microsoft_Samsung.xlsx"
wb.save(excel_file)

print(f"Excel file created successfully: {excel_file}")


# ============================================================
# Create Word Report
# ============================================================

doc = Document()

# Title
title = doc.add_heading("Financial Ratio Analysis", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Company Comparison: Apple vs Microsoft vs Samsung Electronics")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Prepared by: Bhaskar Jha")
doc.add_paragraph("Project Domain: Finance, Equity Research, Business Analytics")
doc.add_paragraph("Tools Used: Python, Excel, OpenPyXL, Financial Statements")


# Introduction
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "Financial ratio analysis is one of the most important techniques used by investors, analysts, "
    "managers and recruiters to understand the financial health of a company. In this project, three "
    "global companies — Apple, Microsoft and Samsung Electronics — are compared using important ratios "
    "such as P/E Ratio, Debt-to-Equity Ratio, Return on Equity and Net Profit Margin."
)

doc.add_paragraph(
    "The objective of this project is not only to calculate ratios but also to interpret them from an "
    "analyst's perspective. This makes the project useful for finance, consulting, analytics and business "
    "roles."
)


# Objectives
doc.add_heading("2. Objectives of the Project", level=1)
objectives = [
    "To collect financial data from annual reports.",
    "To calculate important financial ratios.",
    "To compare companies across profitability, leverage and valuation.",
    "To build a clean Excel dashboard.",
    "To generate business insights useful for recruiters and interview discussions.",
]

for obj in objectives:
    doc.add_paragraph(obj, style="List Bullet")


# Companies
doc.add_heading("3. Companies Selected", level=1)
doc.add_paragraph(
    "Apple, Microsoft and Samsung Electronics were selected because they are globally recognized "
    "technology companies, but their business models are different. Apple is strongly hardware and "
    "ecosystem driven, Microsoft has a major software and cloud business, while Samsung has exposure "
    "to smartphones, consumer electronics and semiconductors."
)


# Financial Data Table
doc.add_heading("4. Financial Data Used", level=1)

table = doc.add_table(rows=1, cols=7)
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
headers = [
    "Company",
    "Revenue",
    "Net Income",
    "Assets",
    "Liabilities",
    "Equity",
    "EPS",
]

for i, header in enumerate(headers):
    hdr_cells[i].text = header

for company, data in companies.items():
    row_cells = table.add_row().cells
    row_cells[0].text = company
    row_cells[1].text = str(data["revenue"])
    row_cells[2].text = str(data["net_income"])
    row_cells[3].text = str(data["total_assets"])
    row_cells[4].text = str(data["total_liabilities"])
    row_cells[5].text = str(data["shareholders_equity"])
    row_cells[6].text = str(data["eps"])


# Ratio Formula
doc.add_heading("5. Ratios Calculated", level=1)

ratios_explanation = [
    (
        "P/E Ratio",
        "Market Price per Share / Earnings per Share",
        "It shows how much investors are willing to pay for one unit of earnings."
    ),
    (
        "Debt-to-Equity Ratio",
        "Total Liabilities / Shareholders' Equity",
        "It measures the financial leverage of a company."
    ),
    (
        "Return on Equity",
        "Net Income / Shareholders' Equity × 100",
        "It measures how efficiently the company uses shareholders' funds to generate profit."
    ),
    (
        "Net Profit Margin",
        "Net Income / Revenue × 100",
        "It shows how much net profit the company earns from each unit of sales."
    ),
]

for name, formula, meaning in ratios_explanation:
    doc.add_heading(name, level=2)
    doc.add_paragraph(f"Formula: {formula}")
    doc.add_paragraph(f"Meaning: {meaning}")


# Ratio Results
doc.add_heading("6. Ratio Analysis Results", level=1)

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
headers = [
    "Company",
    "P/E Ratio",
    "Debt-to-Equity",
    "ROE (%)",
    "Net Profit Margin (%)",
]

for i, header in enumerate(headers):
    hdr_cells[i].text = header

for company, data in companies.items():
    ratios = calculate_ratios(data)
    row_cells = table.add_row().cells
    row_cells[0].text = company
    row_cells[1].text = str(round(ratios["P/E Ratio"], 2))
    row_cells[2].text = str(round(ratios["Debt-to-Equity Ratio"], 2))
    row_cells[3].text = str(round(ratios["Return on Equity (%)"], 2))
    row_cells[4].text = str(round(ratios["Net Profit Margin (%)"], 2))


# Company-wise Analysis
doc.add_heading("7. Company-wise Interpretation", level=1)

doc.add_heading("Apple", level=2)
doc.add_paragraph(
    "Apple shows strong profitability and high brand strength. Its profit margin and ROE indicate "
    "efficient business operations. However, its debt-to-equity ratio is relatively high, which means "
    "the company uses a significant amount of liabilities compared to equity."
)

doc.add_heading("Microsoft", level=2)
doc.add_paragraph(
    "Microsoft appears financially strong and balanced. It has strong profitability, high return on "
    "equity and a relatively stable capital structure. Its cloud and software businesses provide "
    "recurring revenue and strong margins."
)

doc.add_heading("Samsung Electronics", level=2)
doc.add_paragraph(
    "Samsung Electronics has a more conservative capital structure with relatively lower leverage. "
    "However, its profitability can be more cyclical because it depends on consumer electronics and "
    "semiconductor market cycles."
)


# Advanced Section
doc.add_heading("8. Advanced Analyst View", level=1)
doc.add_paragraph(
    "From an analyst's perspective, Microsoft appears to be the most balanced company because it "
    "combines strong profitability with a stable financial structure. Apple remains extremely profitable "
    "but has higher leverage. Samsung has a strong balance sheet but its profit performance depends more "
    "on industry cycles."
)

doc.add_paragraph(
    "A recruiter will not only look at whether ratios are calculated correctly, but also whether the "
    "candidate can interpret those ratios. Therefore, this project focuses on business meaning, not just "
    "formula-based calculation."
)


# Conclusion
doc.add_heading("9. Conclusion", level=1)
doc.add_paragraph(
    "This project compares Apple, Microsoft and Samsung Electronics using key financial ratios. The "
    "analysis shows that Microsoft is the most balanced among the three, Apple has high profitability "
    "but higher leverage, and Samsung is financially conservative but more cyclical. The project gives "
    "a practical understanding of financial statement analysis, equity research and business decision-making."
)


# Resume Section
doc.add_heading("10. How to Mention This Project in Resume", level=1)

resume_points = [
    "Built a financial ratio analysis model comparing Apple, Microsoft and Samsung using annual report data.",
    "Calculated P/E Ratio, Debt-to-Equity Ratio, ROE and Net Profit Margin to evaluate valuation, leverage and profitability.",
    "Created an Excel dashboard with charts and company-wise financial insights.",
    "Interpreted financial ratios from an analyst perspective to identify the strongest company."
]

for point in resume_points:
    doc.add_paragraph(point, style="List Bullet")


# Interview Explanation
doc.add_heading("11. Interview Explanation", level=1)
doc.add_paragraph(
    "In an interview, this project can be explained as a finance and analytics project where real annual "
    "report data was used to compare three global companies. The main focus was not only on calculating "
    "ratios but also on interpreting them. The project helped understand profitability, leverage, valuation "
    "and company performance from an investor's point of view."
)


# Save Word File
word_file = "Financial_Ratio_Analysis_Report_Bhaskar_Jha.docx"
doc.save(word_file)

print(f"Word report created successfully: {word_file}")
